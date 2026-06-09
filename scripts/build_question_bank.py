#!/usr/bin/env python3
"""
build_question_bank.py

Reads scripts/.question_bank_data.json (produced by extract_question_bank.js,
which eval-loads the data constants straight out of app/public/assessment.js) and
renders a clean master question-bank .docx with python-docx.

Output: docs/hive_insightout_question_bank_<YYYYMMDD>.docx
"""
import json
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_PATH = os.path.join(ROOT, "scripts", ".question_bank_data.json")
DATE_YMD = "20260608"          # today's date (2026-06-08)
DATE_HUMAN = "2026-06-08"
OUT_PATH = os.path.join(ROOT, "docs", f"hive_insightout_question_bank_{DATE_YMD}.docx")

blob = json.load(open(DATA_PATH))
D = blob["data"]
PROSE = blob["prose"]
LINES = blob["lines"]
SRC = "assessment.js"

TYPE_NAMES = D["TYPE_NAMES"]

GREY = RGBColor(0x7A, 0x7A, 0x7A)
ACCENT = RGBColor(0x00, 0x8C, 0xBA)
INK = RGBColor(0x1A, 0x2B, 0x33)

counts = {}

# --------------------------------------------------------------------------- #
# helpers
# --------------------------------------------------------------------------- #
doc = Document()
# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK


def src_note(line_key=None, extra=None):
    """Small grey italic one-line source annotation."""
    parts = []
    if line_key is not None:
        ln = LINES.get(line_key)
        if ln:
            parts.append(f"{SRC}:L{ln}")
        else:
            parts.append(SRC)
    if extra:
        parts.append(extra)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("source · " + "  ·  ".join(parts))
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    return p


def qid(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = ACCENT
    return p


def stem(text, label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if label:
        rl = p.add_run(f"[{label}] ")
        rl.bold = True
        rl.font.color.rgb = GREY
        rl.font.size = Pt(9)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def option(label, text=None, indent=0.4):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(1)
    rl = p.add_run(f"{label}" + ("  " if text else ""))
    rl.bold = True
    if text:
        p.add_run(text)
    return p


def body(text, italic=False, indent=0.0, size=10):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    return p


def h1(text):
    doc.add_paragraph()
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = INK
    return h


def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    return h


def bump(section, n=1):
    counts[section] = counts.get(section, 0) + n


# --------------------------------------------------------------------------- #
# TITLE PAGE
# --------------------------------------------------------------------------- #
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(120)
r = t.add_run("InsightOut Assessment\nQuestion Bank")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = INK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run(f"Sourced from assessment.js  ·  {DATE_HUMAN}")
r.font.size = Pt(13)
r.font.color.rgb = ACCENT

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run(
    "Every client-facing question, answer option, and prompt — extracted directly "
    "from the live assessment source (app/public/assessment.js).\n"
    "Each item carries a stable ID and a source line annotation. "
    "Dynamically generated content is noted as such."
)
r.font.size = Pt(10)
r.font.color.rgb = GREY

doc.add_page_break()

# --------------------------------------------------------------------------- #
# STAGE 0
# --------------------------------------------------------------------------- #
h1("Stage 0 — Open-Text Intake")
body("Four free-text prompts. Q3 and Q4 display the client's Q1/Q2 answers back "
     "to them as reference. All four are open text with no answer options.", italic=True)

for i, q in enumerate(D["STAGE0_QUESTIONS"], start=1):
    qid(f"Q0-{i}  ·  {q['title']}")
    stem(q["text"])
    if q.get("showRef"):
        body("Reference box shown: echoes the client's Q1 ('Your words about yourself') "
             "and Q2 ('How others describe you') answers above the input.", italic=True, indent=0.4, size=9)
    src_note(f"stage0:{q['id']}", f"STAGE0_QUESTIONS[{i-1}]")
    bump("Stage 0 prompts")

h2("Tips Interstitial (post-Stage 0)")
body("Shown after Stage 0 Q4, before the Stage 1 sliders ('mid-assessment-reminders' "
     "phase). Doubles as latency cover for the Stage 0 background mini-call.", italic=True)
qid("Q0-TIPS  ·  Mid-Assessment Reminders")
stem(PROSE["midIntro"])
for b in PROSE["midBullets"]:
    option("•", b, indent=0.3)
src_note("prose:mid", "renderMidAssessmentReminders()")
bump("System/interstitial copy")

# --------------------------------------------------------------------------- #
# STAGE 1 — TYPE SLIDERS
# --------------------------------------------------------------------------- #
h1("Stage 1 — Type Sliders")
body("45 statements (9 types × 5 dimensions). Each statement is a single continuous "
     "0–100 slider anchored 'Not like me / Very much like me'. Statements are presented "
     "in the narrative screen order below (two types per screen, paired), with no "
     "type labels visible to the client. The five dimensions per type, in order: "
     "Core motivation (Prioritize) · Focus of attention (Attend) · Resulting "
     "preoccupation (Find) · Energy · Avoidance (Avoid).", italic=True)

order = D["STAGE1_TYPE_SCREEN_ORDER"]      # [3,6,9,1,4,2,8,5,7]
# Pair them two-per-screen exactly as the slider UI does.
screens = [order[i:i+2] for i in range(0, len(order), 2)]
for si, screen_types in enumerate(screens, start=1):
    pairing = " & ".join(f"Type {t} ({TYPE_NAMES[str(t)]})" for t in screen_types)
    h2(f"Slider Screen {si} — {pairing}")
    for t in screen_types:
        ts = str(t)
        body(f"Type {t} — {TYPE_NAMES[ts]}", indent=0.0)
        for s in D["STAGE1_TYPE_STATEMENTS"][ts]:
            qid(f"Q1-T{t}-{s['id'].split('-')[-1]}  ·  {s['id']}")
            stem(s["text"], label=s["dimension"])
            src_note(f"type:{s['id']}", f"STAGE1_TYPE_STATEMENTS[{t}]")
            bump("Stage 1 type statements")

h2("Type Sliders — Optional Follow-Up")
qid("Q1-FT  ·  Type open response")
stem(PROSE["typeOpenPrompt"])
body("Optional free-text. Never gates Continue.", italic=True, indent=0.4, size=9)
src_note("prose:typeOpenPrompt", "renderStage1() — kind:'type-open'")
bump("Stage 1 follow-ups")

# --------------------------------------------------------------------------- #
# STAGE 1 — INSTINCT SLIDERS
# --------------------------------------------------------------------------- #
h1("Stage 1 — Instinct Sliders")
body("15 statements (3 instincts × 5). Same 0–100 'Not like me / Very much like me' "
     "slider format. Presented one instinct block per screen, in order SP → SO → SX.", italic=True)

for inst in D["STAGE1_INSTINCT_ORDER"]:
    full = {"SP": "Self-Preservation", "SO": "Social", "SX": "Sexual / One-to-One"}.get(inst, inst)
    h2(f"Instinct — {inst} ({full})")
    for s in D["STAGE1_INSTINCT_STATEMENTS"][inst]:
        qid(f"Q1-{inst}-{s['id'].split('-')[-1]}  ·  {s['id']}")
        stem(s["text"], label=s["dimension"])
        src_note(f"instinct:{s['id']}", f"STAGE1_INSTINCT_STATEMENTS[{inst}]")
        bump("Stage 1 instinct statements")

h2("Instinct Sliders — Optional Follow-Up")
qid("Q1-FI  ·  Instinct open response")
stem(PROSE["instinctOpenPrompt"])
body("Optional free-text. Never gates Continue.", italic=True, indent=0.4, size=9)
src_note("prose:instinctOpenPrompt", "renderStage1() — kind:'instinct-open'")
bump("Stage 1 follow-ups")

# --------------------------------------------------------------------------- #
# STAGE 2 — CROSS-REFERENCING
# --------------------------------------------------------------------------- #
h1("Stage 2 — Cross-Referencing")
body("Three single-screen questions. Q1 and Q2 are single-select (3 options each); "
     "Q3 is a ranking (rank all three 1st/2nd/3rd). Each maps to a typing framework, "
     "shown as the routing annotation.", italic=True)

routing_map = {
    "Hornevian": "Hornevian — Social Stance (Assertive A / Compliant B / Withdrawn C)",
    "Harmonic": "Harmonic — Conflict Response (Intensity A / Positive B / Competency C)",
    "Centers": "Centers of Intelligence — Decision-Making (Gut / Feelings / Facts)",
}
fw_labels = D["STAGE2_FRAMEWORK_LABELS"]
for i, q in enumerate(D["STAGE2_QUESTIONS"], start=1):
    qid(f"Q2-{i}  ·  {q['title']}")
    stem(q["text"])
    fmt = q.get("format")
    if fmt == "ranking":
        for k in ["a", "b", "c"]:
            mapping = q.get("mapping", {}).get(k, "")
            option(f"{k.upper()}.", f"{q['options'][k]}" + (f"   → {mapping}" if mapping else ""))
        body("Format: ranking — rank each from most important (1st) to least important (3rd).",
             italic=True, indent=0.4, size=9)
    else:
        for k in ["A", "B", "C"]:
            option(f"{k}.", q["options"][k])
        body("Format: single-select.", italic=True, indent=0.4, size=9)
    body(f"Routing: {routing_map.get(q['framework'], q['framework'])}", italic=True, indent=0.4, size=9)
    src_note(f"stage2:{q['id']}", f"STAGE2_QUESTIONS[{i-1}] · framework={q['framework']}")
    bump("Stage 2 questions")

# --------------------------------------------------------------------------- #
# STAGE 3 — PAIRWISE DISCRIMINATION (template)
# --------------------------------------------------------------------------- #
h1("Stage 3 — Pairwise Discrimination (Template)")
body("A discriminating pairwise screen between the AI's top-two candidate types. The "
     "pair shown is chosen at runtime, so this section documents the TEMPLATE plus the "
     "content banks the questions are assembled from.", italic=True)

h2("Shared question template")
qid("Q3-template-1  ·  Q1 Core Motivation (always shown)")
stem(D["STAGE3_Q1_STEM"])
body("Four-option format (same for every pair):", italic=True, indent=0.2, size=9)
for lab in PROSE["fourWayLabels"]:
    option("•", lab, indent=0.4)
body("Person A / Person B carry the two candidate types' core-motivation statements; "
     "the 'Both, but more A/B' options capture a slight lean.", italic=True, indent=0.4, size=9)
src_note("stage3:q1stem", "STAGE3_Q1_STEM + render4WayOptions()")

qid("Q3-template-2  ·  Q2 Avoidance (conditional — high-ambiguity pairs only)")
stem(D["STAGE3_Q2_STEM"])
body("Same four-option (Person A / Person B / Both-more-A / Both-more-B) format. Fires "
     "ONLY when AI Call #1 returns gap='tight' AND the top-two pair is one of the 26 "
     "bespoke pairs below. Otherwise the pair runs Q1 only.", italic=True, indent=0.4, size=9)
src_note("stage3:q2stem", "STAGE3_Q2_STEM")
bump("Stage 3 template stems", 2)

h2("Representative example pair — Type 1 vs Type 6 (bespoke, fires both Q1 & Q2)")
ex = "1-6"
exq2 = D["STAGE3_Q2_PAIRS"][ex]
qid("Example Q1 (core motivation)")
stem(D["STAGE3_Q1_STEM"])
option("Person A", "Type 1: " + D["STAGE3_CORE_MOTIVATIONS"]["1"], indent=0.4)
option("Person B", "Type 6: " + D["STAGE3_CORE_MOTIVATIONS"]["6"], indent=0.4)
option("•", "Both, but more A", indent=0.4)
option("•", "Both, but more B", indent=0.4)
qid("Example Q2 (avoidance) — bespoke copy")
stem(D["STAGE3_Q2_STEM"])
body(f"Label: {exq2['label']}", italic=True, indent=0.4, size=9)
option("Person A", exq2["personA"], indent=0.4)
option("Person B", exq2["personB"], indent=0.4)
option("•", "Both, but more A", indent=0.4)
option("•", "Both, but more B", indent=0.4)
src_note("q2pair:1-6", "STAGE3_Q2_PAIRS['1-6']")

h2("Q1 content bank — 9 Core-Motivation statements (DYNAMICALLY composed)")
body("Q1 pairs are assembled at runtime from these nine statements — any of the 36 "
     "possible type pairs. NOT hardcoded as pairs.", italic=True)
for t in ["1", "2", "3", "4", "5", "6", "7", "8", "9"]:
    qid(f"Q3-CM-{t}  ·  Type {t} ({TYPE_NAMES[t]})")
    stem(D["STAGE3_CORE_MOTIVATIONS"][t])
    src_note(f"motivation:{t}", "STAGE3_CORE_MOTIVATIONS")
    bump("Stage 3 core-motivation statements")

h2("Q2 content bank — 26 bespoke Avoidance pairs (HARDCODED)")
body("Hardcoded, keyed 'lower-higher' type number. Lower-numbered type is Person A. "
     "Each has authored Person A / Person B avoidance copy. Listed by key + label; full "
     "copy is in the source.", italic=True)
for k, pair in D["STAGE3_Q2_PAIRS"].items():
    qid(f"Q3-Q2-{k}  ·  {pair['label']}")
    option("Person A", pair["personA"], indent=0.4)
    option("Person B", pair["personB"], indent=0.4)
    src_note(f"q2pair:{k}", f"STAGE3_Q2_PAIRS['{k}']")
    bump("Stage 3 bespoke Q2 pairs")

h2("Counter-Type comparatives — 5 pairs (HARDCODED)")
body("Hardcoded, keyed by AI Call #1 ct_pair. Counter-type expression is Person A, "
     "lookalike is Person B. One question, same stem as Q1.", italic=True)
for k, ct in D["STAGE3_CT_COMPARATIVES"].items():
    qid(f"Q3-{ct['ctId']}  ·  {ct['label']}")
    stem(D["STAGE3_Q1_STEM"])
    option("Person A", f"(counter-type, Type {ct['counterType']}) " + ct["personA"], indent=0.4)
    option("Person B", f"(lookalike, Type {ct['lookalike']}) " + ct["personB"], indent=0.4)
    src_note(f"ctpair:{k}", f"STAGE3_CT_COMPARATIVES['{k}']")
    bump("Stage 3 CT comparatives")

# --------------------------------------------------------------------------- #
# STAGE 4 — CONFIRMATION (Stress / Security / Habit of Mind)
# --------------------------------------------------------------------------- #
h1("Stage 4 — Confirmation (Stress / Security / Habit of Mind)")
body("Three confirmation instruments — Stress point, Security point, and Habit of "
     "Mind. Like Stage 3, Stage 4 is DYNAMICALLY assembled at runtime: the type(s) and "
     "presentation format are chosen from the AI Call #1 read + the Stage 3 lean. This "
     "section documents the stems, the three presentation formats, and the per-type / "
     "counter-type content banks the questions are assembled from.", italic=True)

h2("Shared question stems")
qid("Q4-STRESS  ·  Stress Point (always shown)")
stem(D["STAGE4_STRESS_STEM"])
src_note("stage4:stressStem", "STAGE4_STRESS_STEM")
qid("Q4-SECURITY  ·  Security Point (always shown)")
stem(D["STAGE4_SECURITY_STEM"])
src_note("stage4:securityStem", "STAGE4_SECURITY_STEM")
qid("Q4-HABIT  ·  Habit of Mind (conditional — see formats)")
stem(D["STAGE4_HABIT_STEM"])
src_note("stage4:habitStem", "STAGE4_HABIT_STEM")
bump("Stage 4 template stems", 3)

h2("Presentation formats (chosen at runtime)")
body("Option A — 3-option (single lead type). The lead type's three options are shown, "
     "SHUFFLED at render. Option index 0 is the canonical / correct answer for that type; "
     "indexes 1 and 2 are distractors drawn from other types' 'energies' (the type's "
     "stress/security line neighbours). Stored as correct / alt1 / alt2.", indent=0.2, size=9.5)
body("Option B — pairwise head-to-head. Person A = the lead type's correct option [0]; "
     "Person B = the second candidate's correct option [0]. Four-way response: Person A / "
     "Person B / Both-but-more-A / Both-but-more-B. (No separate copy — it reuses the "
     "canonical [0] strings from the banks below.)", indent=0.2, size=9.5)
body("Modified Option B — counter-type pairwise. Uses the 5 hardcoded CT comparatives "
     "below (Person A = counter-type, Person B = lookalike). Same four-way response.", indent=0.2, size=9.5)
body("Sequence: Stress then Security always fire. Habit of Mind is appended CONDITIONALLY "
     "(shouldFireHabit): only when Stress and Security disagree, either is unrecognized "
     "(Option A), or any 'slight' answer is given (pairwise modes). It may not appear at "
     "all for a given client.", italic=True, indent=0.2, size=9.5)

instruments = [
    ("Stress", "STAGE4_STRESS", "stress", "Stage 4 stress option sets",
     "Stress — 3-option banks (HARDCODED per type · DYNAMICALLY selected)",
     D["STAGE4_STRESS_STEM"]),
    ("Security", "STAGE4_SECURITY", "security", "Stage 4 security option sets",
     "Security — 3-option banks (HARDCODED per type · DYNAMICALLY selected)",
     D["STAGE4_SECURITY_STEM"]),
    ("Habit", "STAGE4_HABIT", "habit", "Stage 4 habit option sets",
     "Habit of Mind — 3-option banks (HARDCODED per type · DYNAMICALLY selected)",
     D["STAGE4_HABIT_STEM"]),
]
opt_labels = ["Correct (canonical)", "Distractor (alt1)", "Distractor (alt2)"]
for inst_name, const_name, line_prefix, count_key, heading, stem_txt in instruments:
    h2(heading)
    body(f"Stem: {stem_txt}", italic=True, size=9)
    for t in ["1", "2", "3", "4", "5", "6", "7", "8", "9"]:
        idtag = inst_name.upper()
        qid(f"Q4-{idtag}-{t}  ·  Type {t} ({TYPE_NAMES[t]})")
        for oi, txt in enumerate(D[const_name][t]):
            option(opt_labels[oi], txt, indent=0.4)
        src_note(f"{line_prefix}:{t}", const_name)
        bump(count_key)

# Counter-type comparatives (Modified Option B)
ct_idmap = {k: v["ctId"] for k, v in D["STAGE3_CT_COMPARATIVES"].items()}
h2("Counter-Type comparatives — Modified Option B (5 pairs, HARDCODED)")
body("Hardcoded, keyed by CT pair. Fires when Stage 3 counter-type mode produced a "
     "'Both' answer. Person A = counter-type, Person B = lookalike. Each pair carries "
     "its own Stress / Security / Habit copy. Four-way response (Person A / Person B / "
     "Both-more-A / Both-more-B).", italic=True)
for k, ct in D["STAGE4_CT_COMPARATIVE"].items():
    ctid = ct_idmap.get(k, k)
    qid(f"Q4-{ctid}  ·  {ct['label']}")
    for inst in ["stress", "security", "habit"]:
        blk = ct[inst]
        body(inst.capitalize() + ":", indent=0.2, size=9)
        option("Person A", blk["personA"], indent=0.4)
        option("Person B", blk["personB"], indent=0.4)
    src_note(f"s4ct:{k}", f"STAGE4_CT_COMPARATIVE['{k}']")
    bump("Stage 4 CT comparatives")

# --------------------------------------------------------------------------- #
# OPEN-TEXT CLOSES
# --------------------------------------------------------------------------- #
h1("Open-Text Follow-Ups & Closing Question")
body("Optional free-text prompts that appear across the flow. (Stage 0 is entirely "
     "open-text — see Stage 0. Stages 3 and 4 have no free-text close in the code.)", italic=True)

qid("Q1-FT  ·  Stage 1 Type follow-up (repeat)")
stem(PROSE["typeOpenPrompt"])
src_note("prose:typeOpenPrompt", "renderStage1() — kind:'type-open'")

qid("Q1-FI  ·  Stage 1 Instinct follow-up (repeat)")
stem(PROSE["instinctOpenPrompt"])
src_note("prose:instinctOpenPrompt", "renderStage1() — kind:'instinct-open'")

qid("Q-FINAL  ·  Closing open question (before submit)")
stem(PROSE["finalOpenQuestion"])
body(PROSE["finalOpenNote"], italic=True, indent=0.4, size=9)
src_note("prose:finalOpenQuestion", "renderFinalOpen()")
bump("Closing open question")

# --------------------------------------------------------------------------- #
# SYSTEM & FRAMING COPY
# --------------------------------------------------------------------------- #
h1("System State & Framing Copy")
body("Non-question client-facing strings: framing, loading interstitials, completion, "
     "and error states.", italic=True)

h2("Welcome / Landing")
qid("SYS-WELCOME")
stem(PROSE["welcomeHeading"])
for b in PROSE["welcomeBody"]:
    body(b, indent=0.2)
src_note("prose:welcomeHeading", "renderWelcome()")

h2("Loading Interstitials")
qid("SYS-CT  ·  CT mini-call cover (Stage 1 → Stage 2, conditional)")
stem(PROSE["ctHeading"]); body(PROSE["ctSub"], indent=0.2)
src_note("prose:ct", "renderCtAnalyzing()")
qid("SYS-CALL1  ·  AI Call #1 cover (after Stage 2)")
stem(PROSE["call1Heading"]); body(PROSE["call1Sub"], indent=0.2)
body(f"On completion → heading: “{PROSE['call1DoneHeading']}” · sub: “{PROSE['call1DoneSub']}”",
     italic=True, indent=0.2, size=9)
src_note("prose:call1", "renderCall1Analyzing()")
qid("SYS-PROCESSING  ·  AI Call #2 / results generation cover")
stem(PROSE["processingHeading"]); body(PROSE["processingSub"], indent=0.2)
src_note("prose:processing", "renderProcessing()")

h2("Completion Screen")
qid("SYS-CONFIRM")
stem(PROSE["confirmationHeading"])
body(PROSE["confirmationBody1"], indent=0.2)
body(PROSE["confirmationBody2"], indent=0.2)
body("{firstName} and {email} are runtime placeholders.", italic=True, indent=0.2, size=9)
src_note("prose:confirmation", "renderConfirmation()")

h2("Error State")
qid("SYS-ERROR")
stem(PROSE["errorHeading"])
body(PROSE["errorText"], indent=0.2)
body("Button: 'Try again'.", italic=True, indent=0.2, size=9)
src_note("prose:error", "renderError()")
bump("System/framing strings", 6)

# --------------------------------------------------------------------------- #
# DOCUMENT NOTES
# --------------------------------------------------------------------------- #
h1("Document Notes")
notes = [
    "Source of record: app/public/assessment.js on branch `main`. Stage-0/1/2/3 data "
    "was eval-extracted directly from the file's data constants (no transcription); "
    "system/framing prose was lifted verbatim from the render functions and "
    "line-located against the source.",
    "Stage 1 dimension labels: the code's `dimension` field (Core motivation / Focus of "
    "attention / Resulting preoccupation / Energy / Avoidance) is the source of truth. "
    "The brief's shorthand (Prioritize / Attend / Find / Energy / Avoid) maps to these "
    "1:1 and matches the statements' opening verbs.",
    "Stage 1 type-statement IDs are NOT a clean 1–5 run: the middle dimension splits "
    "into a 2a/2b pair (e.g. S3-1, S3-2a, S3-2b, S3-3, S3-4). Still exactly 5 sliders "
    "per type. Stable IDs above use the source ID suffix.",
    "Stage 1 slider screen order is the narrative pairing [3,6][9,1][4,2][8,5][7] "
    "(STAGE1_TYPE_SCREEN_ORDER). Scoring is order-independent; the order is presentation "
    "only. The client sees NO type/instinct labels — those headings are beta-only "
    "(hidden via CSS) and are included here for the bank's clarity.",
    "There is a legacy STAGE1_QUESTIONS array (forced-rank v1 Stage 1, ids q1…q12) still "
    "present near the top of assessment.js. It is NOT used by the v2 slider flow and was "
    "deliberately EXCLUDED from this bank. Flag for cleanup.",
    "Stage 3 is fully template-driven. Q1 (core motivation) pairs are DYNAMICALLY "
    "composed from the 9 STAGE3_CORE_MOTIVATIONS (any of 36 pairs). Q2 (avoidance) is "
    "HARDCODED as 26 bespoke pairs and fires only on gap='tight' + a bespoke top-two. "
    "Counter-type mode uses 5 HARDCODED comparatives. So a pair shows 1 question "
    "(Q1 only) or 2 (Q1 + Q2) — never the 'Q3-1 through Q3-4' the brief anticipated; "
    "the real maximum is two questions per pair.",
    "Stage 4 is the MOST dynamic stage. Three instruments (Stress / Security / Habit) "
    "are each rendered in one of three runtime-selected formats — Option A (3-option, "
    "single lead type), Option B (pairwise head-to-head), or Modified Option B "
    "(counter-type pairwise). Nothing in Stage 4 is a fixed, always-identical question; "
    "the per-type banks below are the raw material the engine assembles from.",
    "Stage 4 Option A (3-option) answers have a FIXED correct answer at index 0; the two "
    "distractors are other types' 'energies' (the type's stress/security line neighbours) "
    "and are SHUFFLED at render, so on-screen option order is non-deterministic. Labels "
    "in this bank (Correct / alt1 / alt2) reflect the source array index, not screen order.",
    "Stage 4 Option B (pairwise) does NOT have its own copy — Person A / Person B reuse "
    "the canonical index-[0] strings from the Stress/Security/Habit banks, shown "
    "head-to-head between the lead and second-candidate types.",
    "Stage 4 Habit of Mind fires CONDITIONALLY (shouldFireHabit) — only on a Stress/"
    "Security disagreement, an unrecognized Option-A answer, or any 'slight' pairwise "
    "answer. For a clean confirmation it never appears.",
    "Minor inconsistency: STAGE4_CT_COMPARATIVE entries carry a `label` but NO `ctId` "
    "field (Stage 3's STAGE3_CT_COMPARATIVES does carry ctId). The 5 keys match Stage 3 "
    "exactly (SO-7, SX-6, SP-3, SP-4, SX-1), so this bank reuses Stage 3's CT-1…CT-5 ids "
    "for stable Q4-CT-n labelling.",
    "SUBTYPE_NAMES (27 instinctual-subtype labels, e.g. 'sp-7' → 'The Epicure') sits just "
    "after the Stage 4 banks. It is report-rendering labelling, not a client-facing "
    "question, and was deliberately EXCLUDED from this bank.",
    "Placeholders {firstName} / {email} on the completion screen are interpolated at "
    "runtime from intake; they appear literally in the source template.",
    "Intake form field labels (First/Last Name, Email, Organization, Coach select) are "
    "form copy rather than questions and were not catalogued as questions; they live in "
    "renderIntake() if needed.",
]
for n in notes:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(n).font.size = Pt(9.5)

# --------------------------------------------------------------------------- #
# SAVE + SANITY COUNT
# --------------------------------------------------------------------------- #
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)

print("Saved:", OUT_PATH)
print("\n=== EXTRACTION COUNT (sanity check) ===")
total = 0
for k in sorted(counts):
    print(f"  {k:38s}: {counts[k]}")
    total += counts[k]
print(f"  {'-'*38}")
print(f"  {'TOTAL items catalogued':38s}: {total}")
